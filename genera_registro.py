genera_registro.py
--- generate_leg_sgi_01.py (原始)


+++ generate_leg_sgi_01.py (修改后)
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script per la generazione del documento Word professionale:
"REGISTRO REQUISITI LEGALI E ALTRI REQUISITI" - Tresun S.r.l.
Conforme agli standard ISO 9001, ISO 14001, ISO 45001
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.table import Table
from docx.text.paragraph import Paragraph


def create_color_rgb(hex_color):
    """Crea un elemento colore RGB per le celle delle tabelle."""
    color = hex_color.replace('#', '')
    return color


def set_cell_shading(cell, hex_color):
    """Imposta il colore di sfondo di una cella."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), create_color_rgb(hex_color))
    tcPr.append(shd)


def set_table_border(table, border_size=4, border_color='404040'):
    """Imposta i bordi esterni della tabella."""
    tbl = table._tbl

    # Get or create tblPr
    tblPr = tbl.tblPr if hasattr(tbl, 'tblPr') and tbl.tblPr is not None else None
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Check if tblBorders already exists
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        # Clear existing borders
        tblBorders.clear()

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tblBorders.append(border)


def add_header_footer(doc):
    """Aggiunge header e footer professionali a tutte le sezioni."""
    for section in doc.sections:
        # Header
        header = section.header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.clear()

        # Text left
        run_left = header_paragraph.add_run('TRESUN S.r.l.')
        run_left.bold = True
        run_left.font.size = Pt(10)
        run_left.font.name = 'Arial'

        # Tab stop for right alignment
        tab_stops = header_paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(16))

        # Tab character
        header_paragraph.add_run('\t')

        # Text right
        run_right = header_paragraph.add_run('REGISTRO REQUISITI LEGALI')
        run_right.font.size = Pt(10)
        run_right.font.name = 'Arial'

        # Add border under header
        header_paragraph.paragraph_format.border_bottom = True
        header_paragraph.paragraph_format.border_bottom_width = Pt(0.5)
        header_paragraph.paragraph_format.border_bottom_color = '003366'

        # Footer
        footer = section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.clear()

        # Create table in footer for layout
        footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        footer_table.autofit = False

        # Left cell - ISO certifications
        left_cell = footer_table.cell(0, 0)
        left_cell.width = Inches(2.5)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_iso = left_para.add_run('ISO 9001 | ISO 14001 | ISO 45001')
        run_iso.font.size = Pt(8)
        run_iso.font.name = 'Arial'

        # Center cell - Page number
        center_cell = footer_table.cell(0, 1)
        center_cell.width = Inches(2)
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_page = center_para.add_run('Pag. ')
        run_page.font.size = Pt(9)
        run_page.font.name = 'Arial'

        # Add page number field
        r = center_para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        r._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r._r.append(fldChar2)

        run_page2 = center_para.add_run(' di ')
        run_page2.font.size = Pt(9)
        run_page2.font.name = 'Arial'

        # Add total pages field
        r2 = center_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        r2._r.append(fldChar3)

        instrText2 = OxmlElement('w:instrText')
        instrText2.text = ' NUMPAGES '
        r2._r.append(instrText2)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r2._r.append(fldChar4)

        # Right cell - Revision
        right_cell = footer_table.cell(0, 2)
        right_cell.width = Inches(2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_rev = right_para.add_run('Rev. 01')
        run_rev.font.size = Pt(9)
        run_rev.font.name = 'Arial'

        # Remove default paragraph borders from footer table
        for row in footer_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def setup_document(doc):
    """Configura le impostazioni base del documento."""
    # Page setup
    for section in doc.sections:
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.orientation = WD_ORIENT.PORTRAIT

    # Clear default content
    if doc.paragraphs:
        doc.paragraphs[0].clear()


def create_cover_page(doc):
    """Crea la pagina di copertina."""
    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Logo text
    logo_run = title_para.add_run('TRESUN S.r.l.\n')
    logo_run.font.name = 'Arial'
    logo_run.font.size = Pt(24)
    logo_run.bold = True
    logo_run.font.color.rgb = None  # Will be set via shading workaround

    # Document title
    doc_title = title_para.add_run('\nREGISTRO REQUISITI LEGALI\nE ALTRI REQUISITI')
    doc_title.font.name = 'Arial'
    doc_title.font.size = Pt(18)
    doc_title.bold = True

    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Document info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.width = Cm(16)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    info_table.columns[0].width = Cm(6)
    info_table.columns[1].width = Cm(10)

    # Header row
    header_data = [
        ('Codice Documento:', 'LEG-SGI-01'),
        ('Revisione:', '01'),
        ('Data Emissione:', datetime.now().strftime('%d/%m/%Y')),
        ('Stato:', 'APPROVATO')
    ]

    for i, (label, value) in enumerate(header_data):
        label_cell = info_table.cell(i, 0)
        value_cell = info_table.cell(i, 1)

        label_para = label_cell.paragraphs[0]
        label_para.clear()
        label_run = label_para.add_run(label)
        label_run.font.name = 'Arial'
        label_run.font.size = Pt(11)
        label_run.bold = True

        value_para = value_cell.paragraphs[0]
        value_para.clear()
        value_run = value_para.add_run(value)
        value_run.font.name = 'Arial'
        value_run.font.size = Pt(11)

        # Apply styling
        set_cell_shading(label_cell, '#F2F2F2')

    # Set table borders
    set_table_border(info_table)

    doc.add_paragraph()

    # Approval matrix
    approval_para = doc.add_paragraph()
    approval_run = approval_para.add_run('MATRICE DI APPROVAZIONE')
    approval_run.font.name = 'Arial'
    approval_run.font.size = Pt(14)
    approval_run.bold = True

    doc.add_paragraph()

    approval_table = doc.add_table(rows=4, cols=4)
    approval_table.width = Cm(18)
    approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    headers = ['Funzione', 'Nome', 'Firma', 'Data']
    roles = [
        ('Redatto da', 'Responsabile SGI', '', ''),
        ('Verificato da', 'Responsabile Qualità', '', ''),
        ('Approvato da', 'Direttore Generale', '', ''),
    ]

    # Header row with dark blue background
    for i, header in enumerate(headers):
        cell = approval_table.cell(0, i)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = None
        set_cell_shading(cell, '#003366')

    # Data rows
    for i, (role, name, signature, date) in enumerate(roles):
        role_cell = approval_table.cell(i+1, 0)
        name_cell = approval_table.cell(i+1, 1)
        sig_cell = approval_table.cell(i+1, 2)
        date_cell = approval_table.cell(i+1, 3)

        for cell, text in [(role_cell, role), (name_cell, name), (sig_cell, ''), (date_cell, '')]:
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            set_cell_shading(cell, '#F2F2F2' if text else '#FFFFFF')

    set_table_border(approval_table)

    # Page break after cover
    doc.add_page_break()


def create_revision_history(doc):
    """Crea la pagina dello storico revisioni."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('STORICO REVISIONI')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Revision table
    rev_table = doc.add_table(rows=3, cols=5)
    rev_table.width = Cm(18)

    # Headers
    headers = ['Rev.', 'Data', 'Descrizione Modifica', 'Redatto', 'Approvato']

    for i, header in enumerate(headers):
        cell = rev_table.cell(0, i)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        set_cell_shading(cell, '#003366')

    # Data rows
    revisions = [
        ('00', '01/01/2024', 'Emissione iniziale', 'RSGI', 'DG'),
        ('01', datetime.now().strftime('%d/%m/%Y'), 'Aggiornamento requisiti legali', 'RSGI', 'DG'),
    ]

    for i, (rev, date, desc, redatto, approvato) in enumerate(revisions):
        cells = [rev_table.cell(i+1, j) for j in range(5)]
        texts = [rev, date, desc, redatto, approvato]

        for cell, text in zip(cells, texts):
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            set_cell_shading(cell, '#F2F2F2')

    set_table_border(rev_table)

    doc.add_page_break()


def create_table_of_contents(doc):
    """Crea l'indice dei contenuti."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('INDICE DEI CONTENUTI')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Contents
    contents = [
        ('1.', 'SCOPO E CAMPO DI APPLICAZIONE'),
        ('2.', 'RIFERIMENTI NORMATIVI'),
        ('3.', 'TERMINI E DEFINIZIONI'),
        ('4.', 'RESPONSABILITÀ E AUTORITÀ'),
        ('5.', 'MODALITÀ OPERATIVE'),
        ('5.1', 'Identificazione dei Requisiti Legali'),
        ('5.2', 'Valutazione della Conformità'),
        ('5.3', 'Aggiornamento e Monitoraggio'),
        ('6.', 'REGISTRO REQUISITI LEGALI'),
        ('7.', 'DOCUMENTI CORRELATI'),
        ('8.', 'ALLEGATI'),
    ]

    for num, title in contents:
        para = doc.add_paragraph()
        num_run = para.add_run(f'{num}  ')
        num_run.font.name = 'Arial'
        num_run.font.size = Pt(11)
        num_run.bold = True

        title_run = para.add_run(title)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(11)

    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    """Aggiunge un'intestazione di sezione con stile personalizzato."""
    para = doc.add_paragraph()

    if level == 1:
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True
    else:
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.space_before = Pt(6)

    return para


def add_normal_paragraph(doc, text):
    """Aggiunge un paragrafo di testo normale."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_after = Pt(6)
    return para


def add_bullet_point(doc, text):
    """Aggiunge un punto elenco."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run('• ' + text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_after = Pt(6)
    return para


def create_styled_table(doc, headers, data, title=None):
    """Crea una tabella con stile professionale."""
    if title:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(11)
        title_run.bold = True
        title_para.paragraph_format.space_after = Pt(6)

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = None
    table.autofit = False

    # Set column widths proportionally
    col_width = Cm(18 / len(headers))
    for col in table.columns:
        col.width = col_width

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        set_cell_shading(cell, '#003366')

    # Data rows
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            cell = row.cells[i]
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(str(cell_data))
            run.font.name = 'Arial'
            run.font.size = Pt(10)

            # Alternate row coloring
            if row_idx % 2 == 0:
                set_cell_shading(cell, '#F2F2F2')
            else:
                set_cell_shading(cell, '#FFFFFF')

    # Set borders
    set_table_border(table, border_size=4, border_color='404040')

    # Add space after table
    doc.add_paragraph()

    return table


def create_main_content(doc):
    """Crea tutto il contenuto principale del documento."""

    # Sezione 1
    add_section_heading(doc, '1. SCOPO E CAMPO DI APPLICAZIONE', level=1)

    add_normal_paragraph(doc,
        'La presente procedura definisce le modalità per l\'identificazione, l\'accesso, '
        'la valutazione e l\'aggiornamento dei requisiti legali e altri requisiti applicabili '
        'al Sistema di Gestione Integrato (Qualità, Ambiente, Salute e Sicurezza sul Lavoro) '
        'di Tresun S.r.l.')

    add_normal_paragraph(doc, 'Il campo di applicazione include:')
    add_bullet_point(doc, 'Tutte le attività svolte presso la sede operativa di Tresun S.r.l.')
    add_bullet_point(doc, 'Tutti i processi aziendali rilevanti ai fini della conformità normativa')
    add_bullet_point(doc, 'Tutto il personale dipendente e collaboratore')

    # Tabella 1
    create_styled_table(
        doc,
        ['Codice', 'Attività', 'Requisiti Applicabili'],
        [
            ['A01', 'Progettazione', 'ISO 9001, D.Lgs. 81/08'],
            ['A02', 'Produzione', 'ISO 14001, D.Lgs. 152/06'],
            ['A03', 'Servizi', 'ISO 45001, Norme contrattuali'],
        ],
        'Tabella 1 - Attività e Requisiti Applicabili'
    )

    # Sezione 2
    add_section_heading(doc, '2. RIFERIMENTI NORMATIVI', level=1)

    add_normal_paragraph(doc,
        'I principali riferimenti normativi per la gestione dei requisiti legali sono:')

    # Tabella 2
    create_styled_table(
        doc,
        ['Norma', 'Titolo', 'Edizione'],
        [
            ['ISO 9001:2015', 'Sistemi di gestione per la qualità', '2015'],
            ['ISO 14001:2015', 'Sistemi di gestione ambientale', '2015'],
            ['ISO 45001:2018', 'Salute e sicurezza sul lavoro', '2018'],
            ['D.Lgs. 81/08', 'Testo Unico sulla Sicurezza', 'Ultima mod. 2023'],
            ['D.Lgs. 152/06', 'Testo Unico Ambientale', 'Ultima mod. 2023'],
        ],
        'Tabella 2 - Riferimenti Normativi Principali'
    )

    # Sezione 3
    add_section_heading(doc, '3. TERMINI E DEFINIZIONI', level=1)

    add_normal_paragraph(doc,
        'Ai fini della presente procedura si intendono le seguenti definizioni:')

    # Tabella 3
    create_styled_table(
        doc,
        ['Termine', 'Definizione'],
        [
            ['Requisito Legale', 'Obbligo normativo derivante da leggi, regolamenti, autorizzazioni'],
            ['Altro Requisito', 'Obblighi derivanti da contratti, accordi volontari, politiche aziendali'],
            ['Conformità', 'Adempimento di un requisito applicabile'],
            ['Non Conformità', 'Mancato adempimento di un requisito applicabile'],
            ['Valutazione', 'Processo di verifica del grado di conformità'],
        ],
        'Tabella 3 - Termini e Definizioni'
    )

    # Sezione 4
    add_section_heading(doc, '4. RESPONSABILITÀ E AUTORITÀ', level=1)

    add_normal_paragraph(doc,
        'Le responsabilità per la gestione dei requisiti legali sono così definite:')

    # Tabella 4
    create_styled_table(
        doc,
        ['Ruolo', 'Responsabilità', 'Autorità'],
        [
            ['Direttore Generale', 'Approvazione procedure e risorse', 'Decisionale'],
            ['Responsabile SGI', 'Coordinamento e monitoraggio', 'Propositiva'],
            ['Responsabile Qualità', 'Verifica conformità', 'Controllo'],
            ['Responsabile Ambiente', 'Monitoraggio normativo ambientale', 'Controllo'],
            ['RSPP', 'Sorveglianza requisiti sicurezza', 'Consultiva'],
            ['Tutto il Personale', 'Applicazione requisiti', 'Operativa'],
        ],
        'Tabella 4 - Matrice delle Responsabilità'
    )

    # Sezione 5
    add_section_heading(doc, '5. MODALITÀ OPERATIVE', level=1)

    add_section_heading(doc, '5.1 Identificazione dei Requisiti Legali', level=2)

    add_normal_paragraph(doc,
        'L\'identificazione dei requisiti legali avviene attraverso:')
    add_bullet_point(doc, 'Analisi del contesto aziendale e delle parti interessate')
    add_bullet_point(doc, 'Consultazione di banche dati normative specializzate')
    add_bullet_point(doc, 'Collaborazione con consulenti esterni qualificati')
    add_bullet_point(doc, 'Monitoraggio di fonti istituzionali (Gazzetta Ufficiale, siti ministeriali)')

    # Tabella 5
    create_styled_table(
        doc,
        ['Fonte', 'Tipo', 'Frequenza Aggiornamento'],
        [
            ['Gazzetta Ufficiale', 'Legislazione nazionale', 'Settimanale'],
            ['BUR Regionali', 'Legislazione regionale', 'Settimanale'],
            ['Siti Ministeriali', 'Linee guida e circolari', 'Mensile'],
            ['Banche Dati', 'Normativa tecnica', 'Continua'],
            ['Consulenti', 'Interpretazioni', 'Su richiesta'],
        ],
        'Tabella 5 - Fonti di Informazione Normativa'
    )

    add_section_heading(doc, '5.2 Valutazione della Conformità', level=2)

    add_normal_paragraph(doc,
        'La valutazione della conformità viene effettuata:')
    add_bullet_point(doc, 'All\'emissione di nuovi requisiti legali')
    add_bullet_point(doc, 'In occasione di audit interni ed esterni')
    add_bullet_point(doc, 'Periodicamente secondo il programma definito')
    add_bullet_point(doc, 'A seguito di cambiamenti nei processi aziendali')

    # Tabella 6
    create_styled_table(
        doc,
        ['Metodo', 'Strumento', 'Responsabile'],
        [
            ['Checklist', 'Lista di verifica', 'Responsabile SGI'],
            ['Audit', 'Programma audit', 'Responsabile Qualità'],
            ['Autovalutazione', 'Questionari', 'Tutti i Responsabili'],
            ['Indicatori', 'KPI di conformità', 'Responsabile SGI'],
        ],
        'Tabella 6 - Metodi di Valutazione Conformità'
    )

    add_section_heading(doc, '5.3 Aggiornamento e Monitoraggio', level=2)

    add_normal_paragraph(doc,
        'Il monitoraggio dei requisiti legali prevede:')
    add_bullet_point(doc, 'Aggiornamento trimestrale del registro requisiti')
    add_bullet_point(doc, 'Riesame annuale completo della conformità')
    add_bullet_point(doc, 'Comunicazione tempestiva delle variazioni significative')
    add_bullet_point(doc, 'Archiviazione documentale delle evidenze')

    # Tabella 7
    create_styled_table(
        doc,
        ['Attività', 'Frequenza', 'Output'],
        [
            ['Ricerca normativa', 'Continua', 'Elenco novità'],
            ['Valutazione impatto', 'Entro 30 giorni', 'Scheda valutazione'],
            ['Aggiornamento registro', 'Trimestrale', 'Registro aggiornato'],
            ['Riesame direzione', 'Annuale', 'Verbale riesame'],
        ],
        'Tabella 7 - Piano di Monitoraggio'
    )

    # Sezione 6
    add_section_heading(doc, '6. REGISTRO REQUISITI LEGALI', level=1)

    add_normal_paragraph(doc,
        'Il Registro Requisiti Legali costituisce il documento principale per la gestione '
        'della conformità normativa. Di seguito sono riportate le tabelle complete.')

    # Tabelle 8-16 con diversi ambiti
    # Tabella 8 - Requisiti Qualità
    create_styled_table(
        doc,
        ['ID', 'Requisito', 'Riferimento', 'Stato', 'Scadenza'],
        [
            ['Q01', 'Manuale Qualità', 'ISO 9001 §4.3', 'Conforme', 'N/A'],
            ['Q02', 'Procedure Documentate', 'ISO 9001 §7.5', 'Conforme', 'N/A'],
            ['Q03', 'Controllo Documenti', 'ISO 9001 §7.5.3', 'Conforme', 'Continuo'],
            ['Q04', 'Gestione Non Conformità', 'ISO 9001 §10.2', 'Conforme', 'Continuo'],
            ['Q05', 'Azioni Correttive', 'ISO 9001 §10.2', 'Conforme', 'Continuo'],
        ],
        'Tabella 8 - Requisiti Sistema Qualità'
    )

    # Tabella 9
    create_styled_table(
        doc,
        ['ID', 'Aspectto Ambientale', 'Normativa', 'Conformità', 'Note'],
        [
            ['A01', 'Gestione Rifiuti', 'D.Lgs. 152/06', 'Conforme', 'Registri aggiornati'],
            ['A02', 'Emissioni Atmosferiche', 'D.Lgs. 152/06', 'Conforme', 'Monitoraggio attivo'],
            ['A03', 'Scarichi Idrici', 'D.Lgs. 152/06', 'Conforme', 'Autorizzazione valida'],
            ['A04', 'Consumo Energetico', 'D.Lgs. 102/14', 'Conforme', 'Diagnosi eseguita'],
        ],
        'Tabella 9 - Requisiti Ambientali'
    )

    # Tabella 10
    create_styled_table(
        doc,
        ['ID', 'Rischio', 'Misura Prevenzione', 'Verifica', 'Stato'],
        [
            ['S01', 'Rischio Chimico', 'DPI specifici', 'Periodica', 'Conforme'],
            ['S02', 'Rischio Rumore', 'Protezioni uditive', 'Fonometria', 'Conforme'],
            ['S03', 'Rischio Vibrazioni', 'Rotazione compiti', 'Valutazione', 'Conforme'],
            ['S04', 'Rischio Ergonomico', 'Postazioni adeguate', 'Sorveglianza', 'Conforme'],
        ],
        'Tabella 10 - Requisiti Sicurezza'
    )

    # Tabella 11
    create_styled_table(
        doc,
        ['ID', 'Obbligo', 'Soggetto', 'Scadenza', 'Stato'],
        [
            ['F01', 'Visita Medica', 'Medico Competente', 'Periodica', 'Programmata'],
            ['F02', 'Formazione', 'Tutto il personale', 'Annuale', 'In corso'],
            ['F03', 'Addestramento', 'Nuovi assunti', 'Prima assegnazione', 'Attivo'],
            ['F04', 'Aggiornamento', 'Preposti/RLS', 'Quinquennale', 'Pianificato'],
        ],
        'Tabella 11 - Formazione e Sorveglianza Sanitaria'
    )

    # Tabella 12
    create_styled_table(
        doc,
        ['ID', 'Documento', 'Emissione', 'Revisione', 'Responsabile'],
        [
            ['D01', 'POS', 'Inizio lavori', 'Per variante', 'Capocantiere'],
            ['D02', 'DUVRI', 'Contratto', 'Annuale', 'Committente'],
            ['D03', 'Piano Emergenza', 'Assunzione', 'Annuale', 'Datore di Lavoro'],
            ['D04', 'Registro Infortuni', 'Evento', 'Immediata', 'Amministrazione'],
        ],
        'Tabella 12 - Documentazione Obbligatoria'
    )

    # Tabella 13
    create_styled_table(
        doc,
        ['ID', 'Autorizzazione', 'Ente', 'Validità', 'Stato'],
        [
            ['AUT01', 'Licenza Edilizia', 'Comune', 'Permanente', 'Valida'],
            ['AUT02', 'AIA/AA', 'Provincia', '5 anni', 'In validità'],
            ['AUT03', 'Scarico Acque', 'Ente Locale', '4 anni', 'Da rinnovare'],
            ['AUT04', 'Gestione Rifiuti', 'Provincia', '5 anni', 'Valida'],
        ],
        'Tabella 13 - Autorizzazioni e Permessi'
    )

    # Tabella 14
    create_styled_table(
        doc,
        ['ID', 'Parte Interessata', 'Requisito', 'Modalità', 'Freq.'],
        [
            ['PI01', 'Cliente', 'Specifiche prodotto', 'Contratto', 'Per ordine'],
            ['PI02', 'Fornitore', 'Condizioni fornitura', 'Accordo', 'Annuale'],
            ['PI03', 'Ente Controllo', 'Adempimenti', 'Comunicazione', 'Come richiesto'],
            ['PI04', 'Comunità Locale', 'Impatto ambientale', 'Report', 'Annuale'],
        ],
        'Tabella 14 - Altri Requisiti Parti Interessate'
    )

    # Tabella 15
    create_styled_table(
        doc,
        ['ID', 'Indicatore', 'Target', 'Risultato', 'Trend'],
        [
            ['KPI01', '% Conformità Legale', '100%', '98%', '↗'],
            ['KPI02', 'Tempo Aggiornamento', '<30 gg', '25 gg', '→'],
            ['KPI03', 'NC Chiuse', '>95%', '97%', '↗'],
            ['KPI04', 'Audit Superati', '100%', '100%', '→'],
        ],
        'Tabella 15 - Indicatori di Performance'
    )

    # Tabella 16
    create_styled_table(
        doc,
        ['ID', 'Azione', 'Priorità', 'Responsabile', 'Scadenza'],
        [
            ['AZ01', 'Agg. Procedura', 'Alta', 'RSGI', '30/06/2024'],
            ['AZ02', 'Formazione', 'Media', 'HR', '30/09/2024'],
            ['AZ03', 'Audit Interno', 'Alta', 'Qualità', '31/10/2024'],
            ['AZ04', 'Riesame', 'Media', 'DG', '15/12/2024'],
        ],
        'Tabella 16 - Piano Azioni Miglioramento'
    )

    # Sezione 7
    add_section_heading(doc, '7. DOCUMENTI CORRELATI', level=1)

    # Tabella documenti correlati
    create_styled_table(
        doc,
        ['Codice', 'Titolo', 'Revisione'],
        [
            ['LEG-SGI-02', 'Procedura Audit Interni', '03'],
            ['LEG-SGI-03', 'Gestione Non Conformità', '02'],
            ['LEG-SGI-04', 'Riesame della Direzione', '01'],
            ['LEG-SGI-05', 'Azioni Correttive', '02'],
            ['FOR-SGI-01', 'Modulo Registro Conformità', '01'],
        ],
        'Documenti Correlati'
    )

    # Sezione 8
    add_section_heading(doc, '8. ALLEGATI', level=1)

    add_normal_paragraph(doc,
        'Di seguito sono elencati gli allegati alla presente procedura:')

    # Allegati
    allegati_data = [
        ['All. A', 'Checklist Identificazione Requisiti'],
        ['All. B', 'Scheda Valutazione Conformità'],
        ['All. C', 'Matrice Applicabilità Normativa'],
        ['All. D', 'Registro Aggiornamenti Normativi'],
        ['All. E', 'Flusso Processuale Gestione Requisiti'],
    ]

    for all_code, all_desc in allegati_
        para = doc.add_paragraph()
        code_run = para.add_run(f'{all_code}: ')
        code_run.font.name = 'Arial'
        code_run.font.size = Pt(11)
        code_run.bold = True
        desc_run = para.add_run(all_desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # Glossario
    add_section_heading(doc, 'GLOSSARIO E ABBREVIAZIONI', level=1)

    glossario_data = [
        ['SGI', 'Sistema di Gestione Integrato'],
        ['RSGI', 'Responsabile Sistema di Gestione Integrato'],
        ['DG', 'Direttore Generale'],
        ['RSPP', 'Responsabile Servizio Prevenzione e Protezione'],
        ['RLS', 'Rappresentante Lavoratori per la Sicurezza'],
        ['MC', 'Medico Competente'],
        ['NC', 'Non Conformità'],
        ['KPI', 'Key Performance Indicator'],
        ['DPI', 'Dispositivi di Protezione Individuale'],
        ['DUVRI', 'Documento Unico Valutazione Rischi Interferenti'],
        ['POS', 'Piano Operativo di Sicurezza'],
        ['AIA', 'Autorizzazione Integrata Ambientale'],
        ['AA', 'Autorizzazione Ambientale'],
    ]

    glossary_table = doc.add_table(rows=len(glossario_data) + 1, cols=2)
    glossary_table.width = Cm(16)

    # Header glossario
    headers_gloss = ['Abbreviazione', 'Significato']
    for i, h in enumerate(headers_gloss):
        cell = glossary_table.cell(0, i)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(h)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        set_cell_shading(cell, '#003366')

    # Data glossario
    for i, (abbr, meaning) in enumerate(glossario_data):
        abbr_cell = glossary_table.cell(i+1, 0)
        mean_cell = glossary_table.cell(i+1, 1)

        for cell, text in [(abbr_cell, abbr), (mean_cell, meaning)]:
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            set_cell_shading(cell, '#F2F2F2' if i % 2 == 0 else '#FFFFFF')

    set_table_border(glossary_table)

    # Final paragraph
    doc.add_paragraph()
    final_para = doc.add_paragraph()
    final_run = final_para.add_run('*** FINE DEL DOCUMENTO ***')
    final_run.font.name = 'Arial'
    final_run.font.size = Pt(11)
    final_run.bold = True
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    """Funzione principale per generare il documento."""
    try:
        print("=" * 60)
        print("Generazione documento Word: REGISTRO REQUISITI LEGALI")
        print("Tresun S.r.l.")
        print("=" * 60)

        # Create document
        print("\n[1/6] Creazione documento...")
        doc = Document()

        # Setup document
        print("[2/6] Configurazione impostazioni pagina...")
        setup_document(doc)

        # Add header and footer
        print("[3/6] Inserimento header e footer...")
        add_header_footer(doc)

        # Create cover page
        print("[4/6] Generazione copertina...")
        create_cover_page(doc)

        # Create revision history
        print("[5/6] Generazione storico revisioni e indice...")
        create_revision_history(doc)
        create_table_of_contents(doc)

        # Create main content
        print("[6/6] Generazione contenuto principale...")
        create_main_content(doc)

        # Save document
        output_filename = 'LEG-SGI-01_Registro_Generato.docx'
        output_path = os.path.join('/workspace', output_filename)

        print(f"\nSalvataggio documento: {output_path}")
        doc.save(output_path)

        # Verify file creation
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print("\n" + "=" * 60)
            print("DOCUMENTO GENERATO CON SUCCESSO!")
            print("=" * 60)
            print(f"Percorso file: {output_path}")
            print(f"Dimensione file: {file_size:,} bytes")
            print(f"Data creazione: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            print("=" * 60)
        else:
            raise FileNotFoundError("Il file non è stato creato correttamente")

        return output_path

    except Exception as e:
        print(f"\nERRORE durante la generazione del documento: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
