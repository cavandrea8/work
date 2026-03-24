# Assicurati di avere python-docx installato:
# pip install python-docx

from docx import Document
from docx.shared import Pt, RGBColor

# Percorsi dei file nella tua repository
file_target = "LEG-SGI-01_Registro_Requisiti_Legali_Tresun_DEFINITIVO_FORMATTATO.docx"
file_source = "Manuale_SGI_Tresun.docx"
file_output = "LEG-SGI-01_Registro_Requisiti_Legali_Tresun_FORMATTATO.docx"

# Carica i documenti
doc_target = Document(file_target)
doc_source = Document(file_source)

# Funzione per copiare stile carattere da un paragrafo all'altro
def copy_style(par_source, par_target):
    if par_source.runs and par_target.runs:
        for run_s, run_t in zip(par_source.runs, par_target.runs):
            run_t.font.name = run_s.font.name
            run_t.font.size = run_s.font.size
            run_t.font.bold = run_s.font.bold
            run_t.font.italic = run_s.font.italic
            run_t.font.underline = run_s.font.underline
            if run_s.font.color.rgb:
                run_t.font.color.rgb = run_s.font.color.rgb

# Applica la formattazione paragrafo per paragrafo
for par_s, par_t in zip(doc_source.paragraphs, doc_target.paragraphs):
    copy_style(par_s, par_t)

# Salva il documento finale con la formattazione aggiornata
doc_target.save(file_output)

print(f"File formattato salvato come '{file_output}'")
